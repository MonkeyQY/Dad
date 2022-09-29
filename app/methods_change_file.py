import io
from importlib.resources import open_binary

import docx

from app import resorces


async def get_file(starting_point: str, end_point: str, date: str) -> io.BytesIO:
    output = create_io()
    doc = preparation_document()
    change_document(doc, date, starting_point, end_point)
    doc.save(output)
    return output


def preparation_document():
    document = open_binary(resorces, 'dad_info.docx')
    doc = docx.Document(document)
    return doc


def create_io() -> io.BytesIO:
    output = io.BytesIO()
    output.name = 'file_dad.docx'
    return output


def change_document(doc, date, starting_point, end_point):
    doc.paragraphs[-2].text = f'Путевой: б/н        от {date}'
    doc.paragraphs[-1].text = 'Маршрут                 '
    doc.paragraphs[-1].add_run(f'{starting_point.title()}-{end_point.title()}').bold = True